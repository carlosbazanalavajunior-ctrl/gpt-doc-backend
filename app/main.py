from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from typing import List, Optional, Union, Any, Tuple
from pathlib import Path
from tempfile import NamedTemporaryFile
from io import BytesIO
import base64
import binascii
import os
import re
import requests
import time
import logging

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


app = FastAPI(title="GPT DOC Backend", version="1.1.0 action-endpoint")


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("docx-timing")


# =========================
# MODELOS
# =========================

class FigureBlock(BaseModel):
    title: Optional[str] = None
    caption: Optional[str] = None
    image_url: Optional[str] = None
    image_base64: Optional[str] = None
    width_inches: Optional[float] = 5.8
    alignment: Optional[str] = "center"


class TableBlock(BaseModel):
    title: Optional[str] = None
    headers: List[str] = Field(default_factory=list)
    rows: List[List[Any]] = Field(default_factory=list)
    data: List[List[Any]] = Field(default_factory=list)
    caption: Optional[str] = None


class SectionBlock(BaseModel):
    heading: Optional[str] = None
    title: Optional[str] = None
    content: List[str] = Field(default_factory=list)
    tables: List[TableBlock] = Field(default_factory=list)
    figures: List[FigureBlock] = Field(default_factory=list)


class GenerateDocumentRequest(BaseModel):
    document_type: str

    year_motto: Optional[str] = ""
    document_code: Optional[str] = ""
    subject: Optional[str] = ""
    city_date: Optional[str] = ""

    # Informe
    addressee: Optional[str] = ""
    reference: Optional[Union[str, List[str]]] = None
    references: Optional[Union[str, List[str]]] = None
    footer_block: Optional[Union[str, List[str]]] = None
    sections: List[SectionBlock] = Field(default_factory=list)

    # Carta
    recipient_name: Optional[str] = ""
    recipient_position: Optional[str] = ""
    recipient_institution: Optional[str] = ""
    greeting: Optional[str] = ""
    body_content: Optional[Union[str, List[str]]] = None
    cc_block: Optional[Union[str, List[str]]] = None


# =========================
# UTILIDADES GENERALES
# =========================

BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = BASE_DIR / "templates"


def normalize_lines(value: Optional[Union[str, List[str]]]) -> List[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    return [str(x) for x in value]


def join_lines(value: Optional[Union[str, List[str]]]) -> str:
    return "\n".join(normalize_lines(value))


def get_reference_text(payload: GenerateDocumentRequest) -> str:
    if payload.reference is not None:
        return join_lines(payload.reference)
    if payload.references is not None:
        return join_lines(payload.references)
    return ""


def alignment_from_text(value: Optional[str]):
    v = (value or "center").strip().lower()
    if v == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if v == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if v == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.CENTER


def insert_paragraph_after(paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def find_paragraph_with_placeholder(document: Document, placeholder: str):
    for p in document.paragraphs:
        if placeholder in p.text:
            return p
    return None


def replace_placeholders_in_paragraph(paragraph, replacements: dict):
    text = paragraph.text
    new_text = text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)
    if new_text != text:
        paragraph.text = new_text


def replace_placeholders_in_document(document: Document, replacements: dict):
    for p in document.paragraphs:
        replace_placeholders_in_paragraph(p, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, replacements)


def clean_empty_paragraphs(document: Document):
    for p in list(document.paragraphs):
        txt = (p.text or "").strip()
        has_drawing = "graphic" in p._element.xml or "drawing" in p._element.xml
        if not txt and not has_drawing:
            try:
                delete_paragraph(p)
            except Exception:
                pass


def choose_template(document_type: str) -> Path:
    document_type = (document_type or "").strip().lower()

    routes = {
        "carta": [
            TEMPLATES_DIR / "carta_profesional.docx",
            TEMPLATES_DIR / "carta_eps_profesional.docx",
            TEMPLATES_DIR / "professional_report_template.docx",
        ],
        "informe": [
            TEMPLATES_DIR / "informe_profesional.docx",
            TEMPLATES_DIR / "informe_eps_profesional.docx",
            TEMPLATES_DIR / "professional_report_template.docx",
        ],
    }

    if document_type not in routes:
        raise HTTPException(status_code=400, detail="document_type debe ser 'carta' o 'informe'.")

    for candidate in routes[document_type]:
        if candidate.exists():
            return candidate

    raise HTTPException(status_code=500, detail="No se encontró ninguna plantilla DOCX válida.")


# =========================
# IMÁGENES
# =========================

def get_image_stream_from_base64(data: str) -> BytesIO:
    if not data:
        raise ValueError("image_base64 está vacío.")

    if "," in data and "base64" in data.split(",", 1)[0].lower():
        data = data.split(",", 1)[1]

    data = re.sub(r"\s+", "", data)

    missing_padding = len(data) % 4
    if missing_padding:
        data += "=" * (4 - missing_padding)

    try:
        binary = base64.b64decode(data, validate=False)
    except binascii.Error as e:
        raise ValueError(f"Base64 inválido: {e}")

    if not binary:
        raise ValueError("El base64 no produjo contenido binario.")

    stream = BytesIO(binary)
    stream.seek(0)
    return stream


def get_image_stream_from_url(url: str) -> BytesIO:
    if not url:
        raise ValueError("image_url está vacío.")

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0 Safari/537.36"
        ),
        "Accept": "image/avif,image/webp,image/apng,image/svg+xml,image/*,*/*;q=0.8",
        "Referer": url,
    }

    response = requests.get(
        url,
        headers=headers,
        timeout=30,
        allow_redirects=True,
        stream=True
    )
    response.raise_for_status()

    content = response.content
    if not content:
        raise ValueError("La URL no devolvió contenido.")

    stream = BytesIO(content)
    stream.seek(0)
    return stream


def get_figure_image_stream(fig: FigureBlock) -> BytesIO:
    if fig.image_base64:
        return get_image_stream_from_base64(fig.image_base64)

    if fig.image_url:
        return get_image_stream_from_url(fig.image_url)

    raise ValueError("La figura no tiene ni image_url ni image_base64.")


def add_image_after(document: Document, paragraph, image_stream: BytesIO, width_inches: float = 5.8, alignment: str = "center"):
    temp_paragraph = document.add_paragraph()
    temp_paragraph.alignment = alignment_from_text(alignment)

    run = temp_paragraph.add_run()
    image_stream.seek(0)
    run.add_picture(image_stream, width=Inches(float(width_inches or 5.8)))

    paragraph._p.addnext(temp_paragraph._p)
    return temp_paragraph


# =========================
# TABLAS
# =========================

def add_table_after(document: Document, paragraph, table_block: TableBlock):
    headers = table_block.headers or []
    rows = table_block.rows or table_block.data or []

    total_cols = 0
    if headers:
        total_cols = len(headers)
    elif rows:
        total_cols = max(len(r) for r in rows if r is not None)

    if total_cols <= 0:
        return paragraph

    table = document.add_table(rows=0, cols=total_cols)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    if headers:
        header_cells = table.add_row().cells
        for i, value in enumerate(headers):
            header_cells[i].text = "" if value is None else str(value)

    for row_data in rows:
        row_cells = table.add_row().cells
        row_values = row_data or []
        for i in range(total_cols):
            value = row_values[i] if i < len(row_values) else ""
            row_cells[i].text = "" if value is None else str(value)

    paragraph._p.addnext(table._tbl)

    last_paragraph = paragraph

    if table_block.caption:
        last_paragraph = insert_paragraph_after(last_paragraph, table_block.caption)
        if last_paragraph.runs:
            for r in last_paragraph.runs:
                r.italic = True
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return last_paragraph


# =========================
# RENDER
# =========================

def add_bold_paragraph_after(paragraph, text: str, align=None):
    new_p = insert_paragraph_after(paragraph)
    run = new_p.add_run(text or "")
    run.bold = True
    if align is not None:
        new_p.alignment = align
    return new_p


def add_normal_paragraph_after(paragraph, text: str, align=None, italic=False):
    new_p = insert_paragraph_after(paragraph)
    run = new_p.add_run(text or "")
    run.italic = italic
    if align is not None:
        new_p.alignment = align
    return new_p


def render_figure(document: Document, current_paragraph, fig: FigureBlock):
    if fig.title:
        current_paragraph = add_bold_paragraph_after(
            current_paragraph,
            fig.title,
            align=alignment_from_text(fig.alignment)
        )

    try:
        image_stream = get_figure_image_stream(fig)

        current_paragraph = add_image_after(
            document=document,
            paragraph=current_paragraph,
            image_stream=image_stream,
            width_inches=float(fig.width_inches or 5.8),
            alignment=fig.alignment or "center"
        )

    except Exception as e:
        warning = f"[No se pudo insertar la imagen: {type(e).__name__}: {str(e)}]"
        current_paragraph = add_normal_paragraph_after(
            current_paragraph,
            warning,
            align=alignment_from_text(fig.alignment),
            italic=True
        )

    if fig.caption:
        current_paragraph = add_normal_paragraph_after(
            current_paragraph,
            fig.caption,
            align=alignment_from_text(fig.alignment),
            italic=True
        )

    return current_paragraph


def render_report_body(document: Document, anchor_paragraph, payload: GenerateDocumentRequest):
    current = anchor_paragraph

    for section in payload.sections:
        heading = section.heading or section.title or ""
        if heading:
            current = add_bold_paragraph_after(current, heading)

        for line in section.content:
            current = add_normal_paragraph_after(current, line)

        for table_block in section.tables:
            if table_block.title:
                current = add_bold_paragraph_after(current, table_block.title, align=WD_ALIGN_PARAGRAPH.CENTER)
            current = add_table_after(document, current, table_block)

        for fig in section.figures:
            current = render_figure(document, current, fig)

    delete_paragraph(anchor_paragraph)


def render_letter_body(document: Document, anchor_paragraph, payload: GenerateDocumentRequest):
    current = anchor_paragraph

    if payload.greeting:
        current = add_normal_paragraph_after(current, payload.greeting)

    for line in normalize_lines(payload.body_content):
        current = add_normal_paragraph_after(current, line)

    for section in payload.sections:
        heading = section.heading or section.title or ""
        if heading:
            current = add_bold_paragraph_after(current, heading)

        for line in section.content:
            current = add_normal_paragraph_after(current, line)

        for table_block in section.tables:
            if table_block.title:
                current = add_bold_paragraph_after(current, table_block.title, align=WD_ALIGN_PARAGRAPH.CENTER)
            current = add_table_after(document, current, table_block)

        for fig in section.figures:
            current = render_figure(document, current, fig)

    delete_paragraph(anchor_paragraph)


# =========================
# GENERACIÓN CENTRAL
# =========================

def build_document_file(payload: GenerateDocumentRequest) -> Tuple[str, str]:
    template_path = choose_template(payload.document_type)
    document = Document(str(template_path))

    body_anchor = find_paragraph_with_placeholder(document, "{{BODY_CONTENT}}")

    common_replacements = {
        "{{YEAR_MOTTO}}": payload.year_motto or "",
        "{{DOCUMENT_CODE}}": payload.document_code or "",
        "{{SUBJECT}}": payload.subject or "",
        "{{REFERENCE_BLOCK}}": get_reference_text(payload),
        "{{CITY_DATE}}": payload.city_date or "",
        "{{FOOTER_BLOCK}}": join_lines(payload.footer_block),
    }

    doc_type = payload.document_type.lower()

    if doc_type == "informe":
        replacements = {
            **common_replacements,
            "{{ADDRESSEE}}": payload.addressee or "",
        }
        replace_placeholders_in_document(document, replacements)

        if body_anchor:
            render_report_body(document, body_anchor, payload)

    elif doc_type == "carta":
        replacements = {
            **common_replacements,
            "{{RECIPIENT_NAME}}": payload.recipient_name or "",
            "{{RECIPIENT_POSITION}}": payload.recipient_position or "",
            "{{RECIPIENT_INSTITUTION}}": payload.recipient_institution or "",
            "{{CC_BLOCK}}": join_lines(payload.cc_block),
        }
        replace_placeholders_in_document(document, replacements)

        if body_anchor:
            render_letter_body(document, body_anchor, payload)

    else:
        raise HTTPException(status_code=400, detail="document_type debe ser 'carta' o 'informe'.")

    clean_empty_paragraphs(document)

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        output_path = tmp.name

    document.save(output_path)
    filename = f"{doc_type}_generado.docx"
    return output_path, filename


# =========================
# ENDPOINTS
# =========================

@app.get("/")
def root():
    return {"message": "GPT DOC Backend operativo"}


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/generate-document")
def generate_document(payload: GenerateDocumentRequest):
    try:
        output_path, filename = build_document_file(payload)

        return FileResponse(
            path=output_path,
            filename=filename,
            media_type=DOCX_MIME
        )

    except Exception as e:
        detail = f"No se pudo generar el documento. Detalle técnico: {type(e).__name__}: {str(e)}"
        raise HTTPException(status_code=500, detail=detail)


@app.post("/generate-document-action")
def generate_document_action(payload: GenerateDocumentRequest):
    output_path = None
    t0 = time.perf_counter()

    try:
        t_build_start = time.perf_counter()
        output_path, filename = build_document_file(payload)
        t_build_end = time.perf_counter()

        t_read_start = time.perf_counter()
        with open(output_path, "rb") as f:
            file_bytes = f.read()
        t_read_end = time.perf_counter()

        t_b64_start = time.perf_counter()
        file_b64 = base64.b64encode(file_bytes).decode("utf-8")
        t_b64_end = time.perf_counter()

        total_time = time.perf_counter() - t0

        logger.info(
            "[DOCX_TIMING] document_type=%s build_s=%.4f read_s=%.4f base64_s=%.4f total_s=%.4f file_size_bytes=%d",
            payload.document_type.lower(),
            t_build_end - t_build_start,
            t_read_end - t_read_start,
            t_b64_end - t_b64_start,
            total_time,
            len(file_bytes),
        )

        return {
            "message": "Documento generado correctamente",
            "document_type": payload.document_type.lower(),
            "filename": filename,
            "timing": {
                "build_s": round(t_build_end - t_build_start, 4),
                "read_s": round(t_read_end - t_read_start, 4),
                "base64_s": round(t_b64_end - t_b64_start, 4),
                "total_s": round(total_time, 4),
                "file_size_bytes": len(file_bytes),
            },
            "openaiFileResponse": [
                {
                    "name": filename,
                    "mime_type": DOCX_MIME,
                    "content": file_b64
                }
            ]
        }

    except Exception as e:
        detail = f"No se pudo generar el documento para action. Detalle técnico: {type(e).__name__}: {str(e)}"
        raise HTTPException(status_code=500, detail=detail)

    finally:
        if output_path and os.path.exists(output_path):
            try:
                os.remove(output_path)
            except Exception:
                pass
